#!/usr/bin/env python3
"""Generate PRICING-PRINCIPLES-AND-SCENARIOS.docx (mirrors the HTML memo)."""
from __future__ import annotations

import os
import sys
import tempfile

_DEPS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_pydeps")
if os.path.isdir(_DEPS):
    sys.path.insert(0, _DEPS)

import pandas as pd

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(ROOT, "PRICING-PRINCIPLES-AND-SCENARIOS.docx")
XLSX_SCENARIOS = os.path.join(ROOT, "pricing-analysis-economics.xlsx")


def _chart_b_series_from_xlsx() -> tuple[list[str], list[float]]:
    """Labels and NOTR % values for Chart B (matches economics_scenarios order)."""
    order = [
        ("BC_baseline_stone", "BC baseline (Stone)"),
        ("segment_upmarket_base", "Upmarket base"),
        ("segment_smb_base", "SMB base"),
        ("portfolio_50_50_merchants", "50/50 SMB+Up"),
    ]
    fallback = (
        ["BC baseline (Stone)", "Upmarket base", "SMB base", "50/50 SMB+Up"],
        [0.77, 1.09, 3.37, 2.23],
    )
    if not os.path.isfile(XLSX_SCENARIOS):
        return fallback
    try:
        df = pd.read_excel(XLSX_SCENARIOS, sheet_name="economics_scenarios")
    except (OSError, ValueError):
        return fallback
    labels: list[str] = []
    vals: list[float] = []
    for sid, lab in order:
        row = df.loc[df["scenario"] == sid]
        if row.empty:
            continue
        ne = row.iloc[0].get("notr_estimated")
        if ne is None or (isinstance(ne, float) and pd.isna(ne)):
            continue
        labels.append(lab)
        vals.append(float(ne) * 100.0)
    if len(vals) < 2:
        return fallback
    return labels, vals


def _bullet(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead) :])
    else:
        p.add_run(text)


def _num_item(doc: Document, lead: str, body: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(lead).bold = True
    p.add_run(" " + body)


def _render_notr_barh_png(
    path: str,
    labels: list[str],
    pct_values: list[float],
    *,
    xlabel: str = "Estimated NOTR (%)",
    xmax_pad: float = 0.15,
) -> None:
    """Write a horizontal bar chart of NOTR percentages to path (PNG)."""
    fig, ax = plt.subplots(figsize=(7.0, max(2.4, 0.45 * len(labels) + 1.2)), dpi=150)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    y = range(len(labels))
    bars = ax.barh(y, pct_values, color="#2E5AAC", height=0.65)
    ax.set_yticks(list(y))
    ax.set_yticklabels(labels, fontsize=10)
    ax.set_xlabel(xlabel, fontsize=10)
    ax.tick_params(axis="x", labelsize=9)
    xmax = max(pct_values) * (1.0 + xmax_pad) if pct_values else 1.0
    ax.set_xlim(0, xmax)
    ax.grid(axis="x", linestyle=":", alpha=0.5)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for bar, v in zip(bars, pct_values):
        ax.text(
            bar.get_width() + xmax * 0.01,
            bar.get_y() + bar.get_height() / 2,
            f"{v:.2f}%",
            va="center",
            ha="left",
            fontsize=9,
        )
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def main() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_paragraph()
    t.add_run("Working draft · PDV payments & pricing").italic = True
    t.runs[0].font.size = Pt(9)

    h = doc.add_heading("Pricing principles, business context & monetization scenarios", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = (
        "This note aligns PDV pricing with the offline channel strategy, the 2026 product vision, "
        "and the PMM diagnosis. It is written for executives: principles first, evidence second, scenarios last. "
        "Quantitative scenarios mirror the pricing deck (https://henryjbpetry.github.io/pos_payments/pricing-deck/) "
        "and the 2026 business case workbook."
    )
    doc.add_paragraph(intro)

    doc.add_heading("Executive summary", level=1)
    _bullet(
        doc,
        "Channel focus: Payments pricing and monetization decisions are optimized for offline GMV and offline revenue through PDV — not for growing online checkout or ecommerce take rate.",
        "Channel focus:",
    )
    _bullet(
        doc,
        "Merchant experience: Fee structures and communication must stay simple and legible — ideally one clear story and one statement across PDV, invoicing, and payments; complexity is a conversion and support tax.",
        "Merchant experience:",
    )
    _bullet(
        doc,
        'Economic story: The same "full stack" logic as the PoS GTM Sales Investment narrative applies (https://gopedroso.github.io/PoS/GTM-Sales-investment/): PoS + offline payments is the bundle that justifies commercial investment; payments are the monetization engine on physical GMV.',
        "Economic story:",
    )
    _bullet(
        doc,
        "Scenarios (illustrative): Base mix 45% crédito à vista / 5% débito / 50% Pix; anticipation fixed to the BC; Upmarket vs SMB list anchors from average(min / max MDR) for crédito and débito (see Chart B and Excel).",
        "Scenarios (illustrative):",
    )

    doc.add_heading("1. Business context", level=1)
    doc.add_heading("1.1 PDV Vision 2026 (multi-store operators)", level=2)
    doc.add_paragraph(
        "The PDV business has demonstrated strong product–market fit (Sean Ellis scores in the high 70s–80% vs a 40% benchmark), "
        "with offline GMV reaching record levels and revenue run-rate scaling with GMV. The strategic shift for 2026 is toward "
        "Multi-Store Operators (MSOs): merchants with multiple locations, staff, and higher offline GMV — a segment that already "
        "represents a large share of POS merchants and offline GMV in the broader market model."
    )
    doc.add_paragraph(
        "Geography (BR / AR / MX): Brazil has overtaken Argentina in recent share in the vision narrative; Mexico and smaller markets "
        "show high offline intensity among adopters. Payments and pricing must work across these realities (regulation, acquirers, and merchant expectations differ)."
    )

    doc.add_heading("1.2 PMM diagnosis (what matters for pricing)", level=2)
    doc.add_paragraph(
        "The PMM diagnostic emphasizes honest perimeter: communicate and price for what the product delivers today, not only the roadmap. "
        "Offline payments — scope, modalities, and narrative — remain partly being mapped; “maquininha” and modality mix must be unpacked before hard commitments."
    )
    doc.add_paragraph("Key PMM themes that intersect pricing:")
    _bullet(doc, "ICP clarity: A large share of trials show weak physical signal; pricing and lifecycle noise follow when the funnel is diluted.", "ICP clarity:")
    _bullet(doc, "Lifecycle: Growth has been largely organic; structured journeys will amplify whatever price model we choose.", "Lifecycle:")
    _bullet(doc, "Pricing risk: March 2026 repricing can pressure conversion — requires disciplined monitoring by segment.", "Pricing risk:")
    _bullet(
        doc,
        "Merchant preference: Evidence suggests heterogeneous appetite (subscription vs per-transaction / CPT-style framing); one-size communication may not fit all ICPs.",
        "Merchant preference:",
    )

    doc.add_heading("1.3 Payments inside the broader story (GTM reference)", level=2)
    doc.add_paragraph(
        "The PoS GTM Sales Investment deck (https://gopedroso.github.io/PoS/GTM-Sales-investment/) frames LatAm opportunity size, segment mix, "
        "and unit economics when PoS and offline payments move together. The executive takeaway for PDV is parallel: payments monetization on offline GMV "
        "is how we capture value from in-store commerce; PoS SaaS alone does not carry the full commercial thesis. Brazil Stone and Argentina dLocal are used "
        "there as illustrative acquirer economics — our PDV pricing work should remain consistent with that “bundle discipline” story even as our actual fee stack differs."
    )

    doc.add_heading("2. How we see payments in the PDV model", level=1)
    _bullet(
        doc,
        "Offline-first revenue: PDV exists to win physical commerce; payment fees should be evaluated against offline GMV and offline revenue, not ecommerce metrics.",
        "Offline-first revenue:",
    )
    _bullet(
        doc,
        'Integrated experience: Merchants already associate PDV with “one place for online + physical”; payments should reinforce that integration without hiding material costs.',
        "Integrated experience:",
    )
    _bullet(
        doc,
        "Transparency beats cleverness: Complex multi-page fee grids erode trust; the PMM doc’s emphasis on clear narrative applies directly to statements, receipts, and sales conversations.",
        "Transparency beats cleverness:",
    )
    _bullet(
        doc,
        "Regulatory and partner reality: MDRs, anticipation, and settlement horizons are set in partnership with acquirers and networks — our principles guide what we optimize and how we explain it, within those constraints.",
        "Regulatory and partner reality:",
    )

    doc.add_heading("3. Pricing principles", level=1)
    _num_item(
        doc,
        "Maximize offline channel results.",
        "Success metrics for this work are offline GMV and offline revenue attributable to PDV-led in-store commerce — not online business outcomes. "
        "Trade-offs that help ecommerce at the expense of clarity or uptake in-store are out of scope.",
    )
    _num_item(
        doc,
        "Radical simplicity for merchants.",
        "Prefer fewer variables, predictable statements, and language that a store manager can explain in one sentence — including how PDV, invoicing, and payments show up on the bill. Avoid opaque bundles that shift risk to the merchant without upside.",
    )
    _num_item(
        doc,
        "Align price communication to ICP reality.",
        "Hybrid vs standalone physical merchants (per PMM hypotheses) may need different framing even if list rates are shared — “simple” does not mean “generic.”",
    )
    _num_item(
        doc,
        "Monitor repricing and conversion jointly.",
        "After material price changes, review conversion and support signals by segment (per PMM risk register) — pricing is a product surface, not only a finance lever.",
    )
    _num_item(
        doc,
        "Bundle economics with PoS.",
        "Strategically, monetize payments in a way that reinforces the full-stack story (PoS + payments) rather than optimizing either line in isolation.",
    )

    doc.add_heading(
        "3.1 Unified billing: invoicing + PDV CPT + Nuvem Pago (working framework)",
        level=2,
    )
    doc.add_paragraph(
        "Today: Invoicing can be sold with PDV and priced per invoice/order, with tiers that follow the online store subscription plan. "
        "That is easy to implement but misaligned with maximizing offline channel GMV and with a single merchant-friendly story."
    )
    doc.add_paragraph(
        "Direction: Move invoicing economics inside the same commercial system as PDV — either the CPT (cost per transaction) meter and/or the PDV subscription (Lite / Plus), "
        "so we can stop charging a separate invoice line for invoicing as an add-on where we choose to bundle."
    )
    doc.add_paragraph("Decision framework (steps):").runs[0].bold = True
    _num_item(
        doc,
        "Anchor the value metric.",
        "Decide what we optimize for (e.g. in-person / POS-attributed activity, Nuvem Pago volume, or unified commerce orders) so invoicing is not priced off ecommerce plan alone.",
    )
    _num_item(
        doc,
        "Path A — Inside CPT.",
        "Treat invoicing as part of the variable stack: same CPT applies to qualifying sales/flows, or CPT bands reflect PDV + invoicing usage. Requires cost-to-serve and metering for invoice events vs POS tickets.",
    )
    _num_item(
        doc,
        "Path B — Inside subscription.",
        "Include invoicing in Lite / Plus (or a single add-on tier) with no per-invoice fee; use fair-use caps or feature gates if needed to protect margin.",
    )
    _num_item(
        doc,
        "Path C — Hybrid.",
        "Small subscription uplift + low/zero marginal on invoices, or blocks of invoices included then variable — only if A/B alone break unit economics.",
    )
    _num_item(
        doc,
        "Pair with payments.",
        "If Nuvem Pago is the strategic monetization engine, align who gets which CPT/subscription treatment when they adopt payments (see CPT options below).",
    )
    _num_item(
        doc,
        "Migration & fairness.",
        "Define grandfathering, credits, and messaging for merchants on legacy per-invoice pricing; avoid surprise bill increases without a clear upgrade story.",
    )
    _num_item(
        doc,
        "Governance.",
        "Finance (transfer pricing), product (metering), legal (commercial terms), support (one statement FAQ).",
    )
    doc.add_paragraph("CPT + payments: keep both designs open (draft):").runs[0].bold = True
    _bullet(
        doc,
        "Option A — CPT waiver / inclusion for PDV when using Nuvem Pago: Simplifies the merchant bill (often no separate invoicing charge); monetization shifts toward payments spread / NOTR. Working preference for simplicity and to stop charging invoicing as a separate meter.",
        "Option A — CPT waiver / inclusion:",
    )
    _bullet(
        doc,
        "Option B — Charge PDV / invoicing separately from payments: Clearer P&L per product line; merchants see distinct software vs payment fees; may feel like a higher total stack price and weaker adoption of Nuvem Pago unless paired with strong value props.",
        "Option B — Separate charges:",
    )
    doc.add_paragraph(
        "Either option must respect eligibility (e.g. merchants who cannot use Nuvem Pago) and internal transfer-pricing rules. Final choice is TBD; this doc keeps both on the table."
    )

    doc.add_heading("4. Monetization scenarios (ultra-executive view)", level=1)
    doc.add_paragraph(
        "The full scenario table lives in the Excel export from the pricing model. For decision-making, two views matter most: "
        "(A) where Brazil Stone sits vs other countries on NOTR in the 2026 BC, and (B) segment economics under a shared payment mix "
        "(45% crédito à vista / 5% débito / 50% Pix), anticipation equal to the BC, and two list-price bases: Upmarket (average of non-promo min MDR crédito & débito − 0.5 p.p., Pix 0.1%, rent in opex) "
        "vs SMB (average of non-promo max MDR crédito & débito − 0.5 p.p., Pix 0.99%, no rent), plus a 50/50 merchant blend (illustrative)."
    )

    doc.add_paragraph("Chart A — NOTR % by provider (2026 BC, “Only CC”)").runs[0].bold = True
    doc.add_paragraph(
        "Shows strategic dispersion: BR Stone vs regional peers in the same business-case export. Not a target to copy — a benchmark map."
    )
    chart_a_labels = ["Getnet", "Stone (BR)", "Stripe", "dLocal QR"]
    chart_a_vals = [0.78, 0.77, 1.96, 0.60]
    temp_pngs: list[str] = []
    try:
        fd_a, path_a = tempfile.mkstemp(suffix=".png", prefix="chart_a_")
        os.close(fd_a)
        temp_pngs.append(path_a)
        _render_notr_barh_png(path_a, chart_a_labels, chart_a_vals, xlabel="NOTR (%)")
        doc.add_picture(path_a, width=Inches(6.0))

        doc.add_paragraph()
        doc.add_paragraph("Chart B — Estimated NOTR by segment (mix + list anchors)").runs[0].bold = True
        doc.add_paragraph(
            "Assumptions: mix 45% crédito à vista / 5% débito / 50% Pix; Pix cost 0.04% of GMV; débito processing cost = crédito BC cost − 0.30 p.p.; "
            "anticipation fee and cost = Stone BC for all. Upmarket: crédito & débito list = average(min MDR crédito, min MDR débito) − 0.5 p.p.; Pix 0.1%; opex includes rent. "
            "SMB: crédito & débito list = average(max MDR crédito, max MDR débito) − 0.5 p.p.; Pix 0.99%; no rent. 50/50 = simple average of Upmarket and SMB NOTR."
        )
        chart_b_labels, chart_b_vals = _chart_b_series_from_xlsx()
        fd_b, path_b = tempfile.mkstemp(suffix=".png", prefix="chart_b_")
        os.close(fd_b)
        temp_pngs.append(path_b)
        _render_notr_barh_png(path_b, chart_b_labels, chart_b_vals, xlabel="Estimated NOTR (%)")
        doc.add_picture(path_b, width=Inches(6.0))
    finally:
        for p in temp_pngs:
            try:
                os.unlink(p)
            except OSError:
                pass

    doc.add_paragraph()
    doc.add_paragraph(
        "Numbers: Chart A from PDV Payments Business Case.xlsx · 2026 BC · NOTR % row. Chart B from economics_scenarios in pricing-analysis-economics.xlsx "
        "(BC_baseline_stone, segment_upmarket_base, segment_smb_base, portfolio_50_50_merchants). "
        "GP and adoption effects are not modeled in these figures — see workbook for linear GP illustrations and limitations."
    )

    doc.add_heading("5. What we still owe (next iterations)", level=1)
    _bullet(doc, "Segment-level price tests aligned to PMM ICP hypotheses (hybrid vs standalone physical).", None)
    _bullet(doc, "Richer anticipation benchmarks as non-promo market data improves.", None)
    _bullet(
        doc,
        "Explicit merchant-facing narrative and fee sheets once modality scope (Pix / credit / NFC, etc.) is locked for launch windows.",
        None,
    )

    doc.add_paragraph()
    src = doc.add_paragraph()
    src.add_run("Sources: ").bold = True
    src.add_run(
        "Internal PDFs PDV Vision 2026: Conquering Multi-Store Operators and [PMM MS] PDV · Diagnóstico e Plano de ação (pos-pricing-model/); "
        "external PoS GTM Sales Investment deck; pricing model outputs in this folder. Working draft — refine with your additional points."
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
